"""
Processadores de arquivos para extrair contexto de diferentes tipos de arquivo
"""
import os
import mimetypes
from typing import Optional, Dict, Any
from django.core.files.uploadedfile import UploadedFile


class BaseFileProcessor:
    """Classe base para processadores de arquivo"""
    
    def __init__(self):
        self.supported_extensions = []
        self.max_file_size = 10 * 1024 * 1024  # 10MB por padrão
    
    def can_process(self, file_path: str) -> bool:
        """Verifica se pode processar o arquivo baseado na extensão"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.supported_extensions
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto do arquivo. Deve ser implementado pelas subclasses."""
        raise NotImplementedError
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Retorna informações básicas do arquivo"""
        stat = os.stat(file_path)
        mime_type, _ = mimetypes.guess_type(file_path)
        
        return {
            'size': stat.st_size,
            'mime_type': mime_type,
            'extension': os.path.splitext(file_path)[1].lower()
        }


class TextFileProcessor(BaseFileProcessor):
    """Processador para arquivos de texto simples"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.txt', '.md', '.csv', '.json', '.html', '.xml']
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto de arquivos texto"""
        try:
            # Tentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        
                    # Se conseguiu ler, processa baseado na extensão
                    ext = os.path.splitext(file_path)[1].lower()
                    
                    if ext == '.json':
                        return self._process_json(content)
                    elif ext == '.csv':
                        return self._process_csv(content)
                    elif ext == '.html' or ext == '.xml':
                        return self._process_html(content)
                    else:
                        return content.strip()
                        
                except UnicodeDecodeError:
                    continue
                    
            return "Erro: Não foi possível decodificar o arquivo"
            
        except Exception as e:
            return f"Erro ao processar arquivo: {str(e)}"
    
    def _process_json(self, content: str) -> str:
        """Formata JSON para melhor legibilidade"""
        try:
            import json
            data = json.loads(content)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except:
            return content
    
    def _process_csv(self, content: str) -> str:
        """Converte CSV em formato legível"""
        try:
            import csv
            import io
            
            # Detectar delimitador
            sample = content[:1024]
            sniffer = csv.Sniffer()
            delimiter = sniffer.sniff(sample).delimiter
            
            # Converter para formato tabular
            reader = csv.reader(io.StringIO(content), delimiter=delimiter)
            rows = list(reader)
            
            if not rows:
                return content
                
            # Formatar como tabela
            formatted_rows = []
            for i, row in enumerate(rows[:100]):  # Limitar a 100 linhas
                if i == 0:
                    formatted_rows.append("Cabeçalhos: " + " | ".join(row))
                    formatted_rows.append("-" * 50)
                else:
                    formatted_rows.append(" | ".join(row))
                    
            if len(rows) > 100:
                formatted_rows.append(f"... e mais {len(rows) - 100} linhas")
                
            return "\n".join(formatted_rows)
            
        except:
            return content
    
    def _process_html(self, content: str) -> str:
        """Extrai texto de HTML"""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remover scripts e estilos
            for script in soup(["script", "style"]):
                script.decompose()
                
            # Extrair texto
            text = soup.get_text()
            
            # Limpar linhas vazias excessivas
            lines = [line.strip() for line in text.splitlines()]
            lines = [line for line in lines if line]
            
            return '\n'.join(lines)
            
        except ImportError:
            # Se beautifulsoup não estiver instalado, retorna conteúdo bruto
            return content
        except:
            return content


class PDFFileProcessor(BaseFileProcessor):
    """Processador para arquivos PDF"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto de PDFs"""
        try:
            # Tentar primeiro com PyPDF4 (mais moderno)
            try:
                import PyPDF4 as PyPDF2
                pdf_module = "PyPDF4"
            except ImportError:
                try:
                    import PyPDF2
                    pdf_module = "PyPDF2"
                except ImportError:
                    return "PyPDF2 ou PyPDF4 não está instalado. Instale com: pip install PyPDF4"
            
            text_content = []
            
            with open(file_path, 'rb') as file:
                try:
                    # Tentar com PdfReader (versão mais nova)
                    if hasattr(PyPDF2, 'PdfReader'):
                        pdf_reader = PyPDF2.PdfReader(file)
                        pages = pdf_reader.pages
                    else:
                        # Fallback para versão antiga
                        pdf_reader = PyPDF2.PdfFileReader(file)
                        pages = [pdf_reader.getPage(i) for i in range(pdf_reader.numPages)]
                    
                    for page_num, page in enumerate(pages):
                        try:
                            # Tentar extract_text() primeiro (versão nova)
                            if hasattr(page, 'extract_text'):
                                page_text = page.extract_text()
                            else:
                                # Fallback para versão antiga
                                page_text = page.extractText()
                                
                            if page_text and page_text.strip():
                                text_content.append(f"--- Página {page_num + 1} ---")
                                text_content.append(page_text.strip())
                                text_content.append("")
                        except Exception as e:
                            text_content.append(f"--- Erro na página {page_num + 1}: {str(e)} ---")
                            
                except Exception as e:
                    return f"Erro ao abrir PDF: {str(e)}"
                        
            if text_content:
                result = '\n'.join(text_content)
                return result if result.strip() else "PDF processado mas não foi possível extrair texto legível"
            else:
                return "Não foi possível extrair texto do PDF"
                
        except Exception as e:
            return f"Erro ao processar PDF: {str(e)}"


class DocxFileProcessor(BaseFileProcessor):
    """Processador para arquivos Word DOCX"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto de arquivos DOCX"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_content = []
            
            # Extrair texto dos parágrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extrair texto de tabelas
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("--- Tabela ---")
                    text_content.extend(table_text)
                    text_content.append("")
            
            return '\n'.join(text_content) if text_content else "Documento vazio"
            
        except ImportError:
            return "python-docx não está instalado. Instale com: pip install python-docx"
        except Exception as e:
            return f"Erro ao processar DOCX: {str(e)}"


class ImageFileProcessor(BaseFileProcessor):
    """Processador para arquivos de imagem"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tiff']
        self.max_file_size = 20 * 1024 * 1024  # 20MB para imagens
    
    def extract_text(self, file_path: str) -> str:
        """
        Para imagens, não extraímos texto diretamente.
        Retornamos metadados básicos da imagem.
        """
        try:
            from PIL import Image
            import os
            
            # Abrir imagem e extrair metadados
            with Image.open(file_path) as img:
                # Informações básicas da imagem
                width, height = img.size
                mode = img.mode
                format_name = img.format or "Unknown"
                
                # Tamanho do arquivo
                file_size = os.path.getsize(file_path)
                size_mb = file_size / (1024 * 1024)
                
                # Construir descrição da imagem
                description = f"""Imagem: {os.path.basename(file_path)}
Formato: {format_name}
Dimensões: {width} x {height} pixels
Modo de cor: {mode}
Tamanho: {size_mb:.2f} MB

Nota: Esta é uma imagem que pode ser enviada diretamente ao usuário via WhatsApp.
Para análise visual do conteúdo da imagem, ela será processada pela IA Vision quando necessário."""
                
                # Verificar se há metadados EXIF (opcional)
                if hasattr(img, '_getexif') and img._getexif():
                    description += "\n\nMetadados EXIF disponíveis."
                
                return description
                
        except ImportError:
            return "PIL (Pillow) não está instalado. Instale com: pip install Pillow"
        except Exception as e:
            return f"Erro ao processar imagem: {str(e)}"


class FileProcessorFactory:
    """Factory para criar processadores baseado no tipo de arquivo"""
    
    def __init__(self):
        self.processors = {
            'text': TextFileProcessor(),
            'pdf': PDFFileProcessor(),
            'docx': DocxFileProcessor(),
            'image': ImageFileProcessor(),
        }
    
    def get_processor(self, file_path: str) -> Optional[BaseFileProcessor]:
        """Retorna o processador apropriado para o arquivo"""
        for processor in self.processors.values():
            if processor.can_process(file_path):
                return processor
        return None
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Processa arquivo e retorna informações + conteúdo extraído"""
        processor = self.get_processor(file_path)
        
        if not processor:
            return {
                'success': False,
                'error': 'Tipo de arquivo não suportado',
                'extracted_text': '',
                'file_info': {}
            }
        
        try:
            # Verificar tamanho do arquivo
            file_info = processor.get_file_info(file_path)
            if file_info['size'] > processor.max_file_size:
                return {
                    'success': False,
                    'error': f'Arquivo muito grande. Máximo: {processor.max_file_size / (1024*1024):.1f}MB',
                    'extracted_text': '',
                    'file_info': file_info
                }
            
            # Extrair texto
            extracted_text = processor.extract_text(file_path)
            
            return {
                'success': True,
                'error': None,
                'extracted_text': extracted_text,
                'file_info': file_info
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Erro durante processamento: {str(e)}',
                'extracted_text': '',
                'file_info': {}
            }
    
    def get_supported_extensions(self) -> list:
        """Retorna lista de extensões suportadas"""
        extensions = []
        for processor in self.processors.values():
            extensions.extend(processor.supported_extensions)
        return sorted(list(set(extensions)))


# Instância global do factory
file_processor = FileProcessorFactory()